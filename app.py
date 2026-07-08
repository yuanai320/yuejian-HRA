# -*- coding: utf-8 -*-
"""
悦检 v3 · Streamlit 交互界面（两段式流程）
=========================================
流程：① 生活方式问卷（前置） → ② HRA 报告解读
     解读入口五种方式：
       A. 粘贴报告文本（PDF/网页复制，或 Excel 中只复制区域列）
       B. 上传 Excel 数据文件（.xls/.xlsx/.csv，惠斯安普 HRA 电子数据整份导入）
       C. 上传 PDF 报告（带文字层可自动提取文本；纯图片 PDF 请用 E）
       D. 上传 Word 报告（.docx，自动提取文本并解析区域值）
       E. 📷 拍照/图片识别（Tesseract 免费 OCR，支持 jpg/png 纸质报告拍照）
依赖：streamlit、pandas；读 .xls 需 xlrd，.xlsx 需 openpyxl，PDF 需 pypdf，Word 需 python-docx。
图片 OCR 需 pytesseract + Pillow + 系统 Tesseract OCR（含中文包 chi_sim，Streamlit Cloud 用 packages.txt 安装）。
无需 API key，知识库内置，离线可演示。
"""

import io
import os
import re
import json
import math
import datetime
import streamlit as st
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
# 中文字体：让雷达图 / 轨迹图的中文标签、标题正常显示（Windows / Mac / Linux 兜底）
plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "PingFang SC",
                                   "Arial Unicode MS", "Noto Sans CJK SC",
                                   "WenQuanYi Micro Hei", "sans-serif"]
plt.rcParams["axes.unicode_minus"] = False  # 正常显示负号

from hra_knowledge import QUESTIONNAIRE, REGIONS, NORMAL_RANGE
from extract import extract_all
from engine_v2 import analyze, to_markdown

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

st.set_page_config(page_title="悦检 · HRA 健康解读", page_icon="❤️", layout="wide")

# ---------- 全局样式：响应式 / 留白 / 去 AI 味 ----------
st.markdown("""
<style>
.block-container{max-width:900px;padding-top:1.8rem;padding-bottom:3rem;}
html,body,[class*="st-"]{font-size:17px;line-height:1.72;}
h1{font-size:2rem;font-weight:700;letter-spacing:-.02em;}
h2{font-size:1.45rem;font-weight:700;margin-top:2rem;margin-bottom:.2rem;}
h3{font-size:1.18rem;font-weight:600;}
p,li{font-size:1.02rem;}
[data-testid="stMetric"]{background:#fff;border:1px solid #eef0f3;border-radius:14px;padding:.9rem 1.1rem;box-shadow:0 1px 3px rgba(16,30,54,.05);}
[data-testid="stMetric"] label{font-size:.9rem;color:#5b6470;font-weight:500;}
.stAlert{border-radius:12px;border-left:4px solid #2e7d32;background:#f4faf5;}
hr{border-color:#eceff3;}
.stButton>button{border-radius:10px;}
@media (max-width:640px){
 .block-container{padding-top:1rem;padding-left:.9rem;padding-right:.9rem;}
 html,body,[class*="st-"]{font-size:16px;line-height:1.66;}
 h1{font-size:1.55rem;} h2{font-size:1.28rem;}
 [data-testid="stMetric"]{padding:.7rem .8rem;}
}
</style>
""", unsafe_allow_html=True)
st.title("❤️ 悦检 · HRA 健康风险评估解读与干预助手")
st.caption("健康服务与管理专业作品 · 科普参考，不替代医生诊断")

# —— 依赖自检：缺包会直接导致上传 PDF / Word / Excel(.xls) 失败 ——
_missing_deps = []
for _m, _pkg in [("pypdf", "pypdf"), ("docx", "python-docx"),
                 ("xlrd", "xlrd"), ("openpyxl", "openpyxl"),
                 ("pytesseract", "pytesseract"), ("PIL", "Pillow")]:
    try:
        __import__(_m)
    except ImportError:
        _missing_deps.append(_pkg)
if _missing_deps:
    st.warning(
        "⚠️ 检测到以下依赖未安装，上传 PDF / Word / Excel(.xls) 会失败：\n\n"
        "```bash\npip install " + " ".join(_missing_deps) + "\n```\n\n"
        "安装后重启本程序（终端先按 `Ctrl+C`，再 `py -m streamlit run app.py`）。\n"
        "👉 想让朋友**免装 Python 直接打开网页**？见页面底部「🚀 免装部署」说明。"
    )

# 会话状态：控制两段式流程
if "step" not in st.session_state:
    st.session_state.step = "问卷"
if "result" not in st.session_state:
    st.session_state.result = None


# ---------------- 工具：把 DataFrame 的一行转成 (profile, regions) ----------------
def df_row_to_data(df, idx):
    """从 pandas DataFrame 的某一行抽取会员信息与区域检测值。"""
    profile, regions = {}, {}
    col_map = {"姓名": "姓名", "性别": "性别", "身高 (cm)": "身高", "身高(cm)": "身高",
               "体重 (kg)": "体重", "体重(kg)": "体重"}
    for col, key in col_map.items():
        if col in df.columns:
            v = df.iloc[idx][col]
            if pd.notna(v):
                s = str(v).strip()
                if key in ("身高", "体重"):
                    m = re.search(r"\d{2,3}", s)
                    if m:
                        profile[key] = int(m.group(0))
                else:
                    profile[key] = s
    # BMI — 从「人体成分分析及建议」列精准提取（防体脂率等先出现时误取）
    for col in df.columns:
        if "BMI" in str(col) or "人体成分" in str(col):
            v = df.iloc[idx][col]
            if pd.notna(v):
                m = re.search(r"BMI[（(]?[^:：]*[:：]?\s*(\d{1,2}(?:\.\d+)?)", str(v), re.I)
                if m:
                    profile["BMI"] = float(m.group(1))
                    break
    # 区域检测值
    for col in df.columns:
        if col in REGIONS:
            v = df.iloc[idx][col]
            if pd.notna(v):
                try:
                    regions[col] = int(round(float(v)))
                except (ValueError, TypeError):
                    pass
    return profile, regions


def read_uploaded_file(uploaded):
    """读取上传的 xls/xlsx/csv → DataFrame。"""
    name = uploaded.name.lower()
    raw = uploaded.getvalue()
    if name.endswith(".csv"):
        return pd.read_csv(io.BytesIO(raw))
    if name.endswith(".xlsx"):
        return pd.read_excel(io.BytesIO(raw), engine="openpyxl")
    # .xls 默认 xlrd
    try:
        return pd.read_excel(io.BytesIO(raw), engine="xlrd")
    except Exception:
        return pd.read_excel(io.BytesIO(raw))


def extract_pdf_text(raw_bytes):
    """从 PDF 提取纯文本（pypdf）。返回 (文本, 错误信息)。"""
    try:
        from pypdf import PdfReader
    except Exception:
        return None, "读取 PDF 需安装 pypdf：`pip install pypdf`"
    try:
        reader = PdfReader(io.BytesIO(raw_bytes))
        texts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(texts), None
    except Exception as e:
        return None, f"PDF 解析失败：{e}"


def extract_docx_text(raw_bytes):
    """从 Word(.docx) 提取纯文本（python-docx，含表格）。返回 (文本, 错误信息)。"""
    try:
        import docx
    except Exception:
        return None, "读取 Word 需安装 python-docx：`pip install python-docx`"
    try:
        doc = docx.Document(io.BytesIO(raw_bytes))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)
        return "\n".join(texts), None
    except Exception as e:
        return None, f"Word 解析失败：{e}"


def ocr_image_to_text(raw_bytes):
    """用 Tesseract 对图片做中英文 OCR，返回 (文本, 错误信息)。无 tesseract 时给出安装提示。"""
    try:
        from PIL import Image, ImageOps
        import pytesseract
    except ImportError as e:
        return None, (f"图片识别需要 Pillow + pytesseract：`pip install Pillow pytesseract`，"
                      f"并安装系统 Tesseract OCR（含中文包 chi_sim）。详情见「🚀 免装部署」。\n原始错误：{e}")
    try:
        img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
        # 提升识别率：转灰度 + 自动对比度增强
        img = ImageOps.grayscale(img)
        img = ImageOps.autocontrast(img)
        # 小字报告适当放大，更易识别
        w, h = img.size
        if max(w, h) < 2000:
            scale = 2000.0 / max(w, h)
            img = img.resize((int(w * scale), int(h * scale)))
        text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        return text, None
    except Exception as e:
        return None, f"图片识别失败：{e}（请确认服务器已安装 Tesseract OCR 及中文包 chi_sim）"


def record_label(df, idx):
    """生成记录下拉选项文案。"""
    parts = []
    for c in ("姓名", "会员号", "检查日期"):
        if c in df.columns and pd.notna(df.iloc[idx][c]):
            parts.append(str(df.iloc[idx][c]).strip())
    if not parts:
        parts.append(f"第{idx+1}条")
    return " ｜ ".join(parts)


def _radar_chart(systems):
    """各系统最大功能偏离雷达图（对标 InsideTracker Body Map，一眼见弱项）。"""
    labels = [s["system"].replace("系统", "") for s in systems]
    vals = [min(s["maxabs"], 100) for s in systems]
    n = len(labels)
    if n < 3:
        return None
    ang = [i / float(n) * 2 * math.pi for i in range(n)] + [0.0]
    vals2 = vals + [vals[0]]
    fig, ax = plt.subplots(figsize=(4.6, 4.6), subplot_kw=dict(polar=True))
    ax.set_theta_offset(math.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(ang[:-1])
    ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylim(0, 100)
    ax.set_yticks([NORMAL_RANGE, 50, 100])
    ax.set_yticklabels([f"正常±{NORMAL_RANGE}", "50", "100"], fontsize=7, color="#888")
    ax.plot(ang, vals2, color="#e74c3c", linewidth=2)
    ax.fill(ang, vals2, color="#e74c3c", alpha=0.22)
    ax.plot(ang, [NORMAL_RANGE] * len(ang), color="#2ecc71", linestyle="--", linewidth=1)
    ax.set_title("各系统最大功能偏离（越靠外偏差越大）", fontsize=10, pad=10)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=110)
    plt.close(fig)
    buf.seek(0)
    return buf


def _draw_trajectory(traj):
    """健康轨迹：多次基线叠加成趋势曲线（长期监测 / 复利）。"""
    traj = sorted(traj, key=lambda x: x.get("date", ""))
    n = len(traj)
    idx = list(range(n))
    dates = [t.get("date", f"#{i+1}") for i, t in enumerate(traj)]
    hi = [t.get("health_index", 0) for t in traj]
    fig, ax = plt.subplots(figsize=(6, 2.8))
    ax.plot(idx, hi, marker="o", color="#2ecc71", linewidth=2)
    ax.set_xticks(idx)
    ax.set_xticklabels(dates, fontsize=8, rotation=30, ha="right")
    ax.set_ylim(0, 100)
    ax.set_ylabel("健康指数")
    ax.set_title("综合健康指数轨迹")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=110)
    plt.close(fig)
    buf.seek(0)
    st.image(buf)
    delta = hi[-1] - hi[0]
    arrow = "↑ 改善" if delta > 0 else ("↓ 下滑" if delta < 0 else "→ 持平")
    st.success(f"首次（{dates[0]}）→ 本次（{dates[-1]}）：健康指数 {hi[0]} → {hi[-1]}（{arrow} {abs(delta)} 分）")
    # 各系统最大偏离复线
    sys_names = [s["system"] for s in traj[-1].get("systems", [])]
    fig2, ax2 = plt.subplots(figsize=(6, 3.4))
    for sname in sys_names:
        ys = [next((s["maxabs"] for s in t.get("systems", []) if s["system"] == sname), None)
              for t in traj]
        if any(v is not None for v in ys):
            ax2.plot(idx, ys, marker=".", label=sname.replace("系统", ""))
    ax2.axhline(NORMAL_RANGE, color="#2ecc71", linestyle="--", linewidth=1, label="正常阈值")
    ax2.set_xticks(idx)
    ax2.set_xticklabels(dates, fontsize=8, rotation=30, ha="right")
    ax2.set_ylabel("最大偏离")
    ax2.set_title("各系统功能偏离轨迹")
    ax2.legend(fontsize=7)
    buf2 = io.BytesIO()
    fig2.savefig(buf2, format="png", bbox_inches="tight", dpi=110)
    plt.close(fig2)
    buf2.seek(0)
    st.image(buf2)


# ---------------- PDF 导出：把报告渲染成带中文的 PDF ----------------
def _register_cjk_font():
    """为 reportlab 注册一个支持中文的 TTF；返回字体名，失败返回 None（退回 Helvetica）。"""
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import matplotlib.font_manager as fm
    except Exception:
        return None
    cands = ["Microsoft YaHei", "SimHei", "PingFang SC", "Arial Unicode MS",
             "Noto Sans CJK SC", "WenQuanYi Micro Hei", "Noto Sans CJK JP",
             "Source Han Sans SC"]
    path = None
    for name in cands:
        try:
            p = fm.findfont(fm.FontProperties(family=name), fallback_to_default=False)
            if p and os.path.exists(p):
                path = p
                break
        except Exception:
            continue
    if not path:
        for f in fm.fontManager.ttflist:
            if any(k in f.name for k in ["CJK", "YaHei", "SimHei", "Hei",
                                         "PingFang", "WenQuanYi", "Han Sans",
                                         "Noto Sans CJK"]):
                if os.path.exists(f.fname):
                    path = f.fname
                    break
    if not path:
        return None
    # reportlab 不直接支持 .ttc，用 fontTools 抽出首个字体转存为临时 .ttf
    if path.lower().endswith(".ttc"):
        try:
            from fontTools.ttLib.ttCollection import TTCollection
            import tempfile as _tempfile
            coll = TTCollection(path)
            tmp = _tempfile.NamedTemporaryFile(suffix=".ttf", delete=False, dir=BASE_DIR)
            coll.fonts[0].save(tmp.name)
            path = tmp.name
        except Exception:
            return None
    try:
        pdfmetrics.registerFont(TTFont("CJK", path))
        # 同时注册粗体别名：避免报告里 **加粗** 在段落/表格中找不到 CJK-Bold 而回退成方块
        try:
            pdfmetrics.registerFont(TTFont("CJK-Bold", path))
        except Exception:
            pass
        return "CJK"
    except Exception:
        return None


def _inline_md(t):
    """把 Markdown 行内的 **粗体** 转成 reportlab 的 <b>，并转义 HTML 特殊字符。"""
    parts = re.split(r"(\*\*.+?\*\*)", t)
    out = []
    for p in parts:
        if len(p) >= 4 and p.startswith("**") and p.endswith("**"):
            inner = p[2:-2].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            out.append("<b>" + inner + "</b>")
        else:
            out.append(p.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    return "".join(out)


def _flush_md_table(rows, style):
    from reportlab.platypus import Table, TableStyle, Spacer, Paragraph
    from reportlab.lib import colors
    # 关键修复：单元格必须用 Paragraph 包裹，才能继承中文字体（CJK）；
    # 直接传字符串会被 reportlab 用默认 Helvetica 渲染，导致整张表中文变方块
    data = [[Paragraph(_inline_md(c), style) for c in r] for r in rows]
    t = Table(data, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cccccc")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#fdecea")),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    return [t, Spacer(1, 4)]


def _md_to_pdf(md_text, radar_buf=None):
    """把 to_markdown 生成的报告渲染为带中文的 PDF 字节（含雷达图）。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image, HRFlowable)
    font = _register_cjk_font() or "Helvetica"
    ss = getSampleStyleSheet()

    def S(name, **kw):
        kw.setdefault("wordWrap", "CJK")
        kw.setdefault("fontName", font)
        return ParagraphStyle(name, parent=ss["Normal"], **kw)

    h2 = S("H2", fontSize=14, spaceBefore=10, spaceAfter=5, textColor=colors.HexColor("#c0392b"), leading=18)
    h3 = S("H3", fontSize=11.5, spaceBefore=7, spaceAfter=3, textColor=colors.HexColor("#2c3e50"), leading=15)
    h4 = S("H4", fontSize=10.5, spaceBefore=5, spaceAfter=2, leading=14)
    body = S("Body", fontSize=9.5, leading=15, spaceAfter=4)
    quote = S("Quote", fontSize=9, leading=14, leftIndent=10, textColor=colors.HexColor("#555"),
              backColor=colors.HexColor("#f5f5f5"), borderPadding=5, spaceAfter=5)
    liststyle = S("List", fontSize=9.5, leading=14, leftIndent=14, spaceAfter=2)

    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, topMargin=16 * mm, bottomMargin=15 * mm,
                            leftMargin=16 * mm, rightMargin=16 * mm, title="HRA综合解读报告")
    flow = [Paragraph("HRA 健康风险评估解读报告", h2),
            HRFlowable(width="100%", thickness=1, color=colors.HexColor("#c0392b")),
            Spacer(1, 6)]
    if radar_buf is not None:
        try:
            radar_buf.seek(0)
            img = Image(radar_buf, width=108 * mm, height=108 * mm)
            img.hAlign = "CENTER"
            flow.append(img)
            flow.append(Spacer(1, 6))
        except Exception:
            pass
    table_buf = []
    for line in md_text.splitlines():
        s = line.rstrip()
        if not s.strip():
            if table_buf:
                flow += _flush_md_table(table_buf, body)
                table_buf = []
            flow.append(Spacer(1, 3))
            continue
        if s.startswith("### "):
            if table_buf:
                flow += _flush_md_table(table_buf, body)
                table_buf = []
            flow.append(Paragraph(_inline_md(s[4:]), h4))
            continue
        if s.startswith("## "):
            if table_buf:
                flow += _flush_md_table(table_buf, body)
                table_buf = []
            flow.append(Paragraph(_inline_md(s[3:]), h3))
            continue
        if s.startswith("# "):
            if table_buf:
                flow += _flush_md_table(table_buf, body)
                table_buf = []
            flow.append(Paragraph(_inline_md(s[2:]), h2))
            continue
        if s.startswith("> "):
            if table_buf:
                flow += _flush_md_table(table_buf, body)
                table_buf = []
            flow.append(Paragraph(_inline_md(s[2:]), quote))
            continue
        if re.match(r"^\s*[-*]\s+", s):
            if table_buf:
                flow += _flush_md_table(table_buf, body)
                table_buf = []
            item = re.sub(r"^\s*[-*]\s+", "", s)
            flow.append(Paragraph("• " + _inline_md(item), liststyle))
            continue
        if s.strip().startswith("|") and s.strip().endswith("|"):
            cells = [c.strip() for c in s.strip().strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-+:?", c) for c in cells):
                continue
            table_buf.append(cells)
            continue
        if table_buf:
            flow += _flush_md_table(table_buf, body)
            table_buf = []
        flow.append(Paragraph(_inline_md(s), body))
    if table_buf:
        flow += _flush_md_table(table_buf, body)
    doc.build(flow)
    out.seek(0)
    return out.getvalue()


# ---------------- 步骤一：生活方式问卷 ----------------
if st.session_state.step == "问卷":
    st.header("📝 第一步：生活方式问卷")
    st.info("先了解你的生活习惯，解读时会结合问卷做个性化建议。填完点「开始解读」。")
    answers = {}
    cols = st.columns(2)
    for i, item in enumerate(QUESTIONNAIRE):
        with cols[i % 2]:
            answers[item["key"]] = st.radio(item["q"], item["options"], key="q_" + item["key"])
    if st.button("✅ 开始解读 →", type="primary", use_container_width=True):
        st.session_state.q = answers
        st.session_state.step = "报告"
        st.session_state.result = None
        st.rerun()

# ---------------- 步骤二：报告解读 ----------------
else:
    st.header("📋 第二步：导入你的 HRA 报告")
    if st.button("← 返回修改问卷"):
        st.session_state.step = "问卷"
        st.session_state.result = None
        st.rerun()

    st.info(f"把 HRA 报告里的文字（含各区域「检测值」，正常值 ±{NORMAL_RANGE}）复制粘贴到下方，"
            f"或直接从 Excel 整份**上传 .xls/.xlsx/.csv** 自动读取；也可以点「填入真实样本」试用真实电子数据。")

    # 合成示例（区域名已对齐真实导出）
    demo = ("|姓名|示例|性别:|女|身高(cm)|165|体重(Kg)|60|\n"
            "|气管附近|气管附近|-66|支气管区域|支气管区域|-66|右肺下叶区域|右肺下叶区域|-50|\n"
            "|胃区域|胃区域|-55|十二指肠区域|十二指肠区域|-51|肝右页|肝右页|-65|胆囊区域|胆囊区域|-59|\n"
            "|左颈动脉|左颈动脉|-53|右颈动脉|右颈动脉|-47|冠状血管|冠状血管|54|心脏区域|心脏区域|-56|\n"
            "|甲状腺区域|甲状腺区域|-47|左杏仁体|左杏仁体|30|C5|-45|Th2|-67|Co1|50|")

    tab_paste, tab_file, tab_photo = st.tabs(["📝 粘贴报告文本", "📎 上传文件(Excel/PDF/Word)", "📷 拍照/图片识别"])

    with tab_paste:
        c1, c2 = st.columns([5, 1])
        with c1:
            text = st.text_area("在此粘贴 HRA 报告文本", height=220, key="report_text")
        with c2:
            if st.button("📥 填入示例"):
                st.session_state.report_text = demo
                st.rerun()
            if st.button("📥 填入真实样本"):
                try:
                    with open(os.path.join(BASE_DIR, "sample_real.tsv"), encoding="utf-8") as f:
                        st.session_state.report_text = f.read()
                    st.rerun()
                except FileNotFoundError:
                    st.error("未找到 sample_real.tsv，请确认文件在同目录。")
        if st.button("✨ 生成综合解读", type="primary", use_container_width=True, key="gen_paste"):
            if not text.strip():
                st.warning("请先粘贴报告文本，或点「填入示例 / 真实样本」。")
            else:
                profile, regions = extract_all(text)
                if not regions:
                    st.error("未能从文本中识别到任何区域检测值。若从 Excel 整行复制，建议改用「上传数据文件」标签页，"
                             "或只复制区域列。也可先点「填入真实样本」看格式。")
                else:
                    res = analyze(regions, profile, st.session_state.get("q", {}))
                    st.session_state.result = res
                    st.success(f"✅ 已识别 {len(regions)} 个区域、{len(res['systems'])} 个系统，"
                               f"主要矛盾：**{res['principal']['system']}**")

    with tab_file:
        uploaded = st.file_uploader("选择 HRA 报告文件（Excel / PDF / Word 均可）",
                                    type=["xls", "xlsx", "csv", "pdf", "docx", "doc"])
        if uploaded:
            name = uploaded.name.lower()
            # —— 文档类（PDF / Word）：提取文本 → 解析 ——
            if name.endswith((".pdf", ".docx", ".doc")):
                if name.endswith(".pdf"):
                    txt, err = extract_pdf_text(uploaded.getvalue())
                elif name.endswith(".docx"):
                    txt, err = extract_docx_text(uploaded.getvalue())
                else:  # .doc 旧版
                    st.warning("暂不支持旧版 .doc 格式，请另存为 .docx，或复制文本到「📝 粘贴报告文本」。")
                    txt, err = None, "skip"
                if err and err != "skip":
                    st.error(err)
                elif txt:
                    st.caption(f"已从文档提取 {len(txt)} 字文本，正在解析区域检测值…")
                    profile, regions = extract_all(txt)
                    if not regions:
                        st.error("未能从文档中识别区域检测值。建议：用 Excel 导出 .xls 上传，"
                                 "或把报告关键文本复制到「📝 粘贴报告文本」。")
                    else:
                        if st.button("✨ 解读该文档", type="primary", use_container_width=True, key="gen_doc"):
                            res = analyze(regions, profile, st.session_state.get("q", {}))
                            st.session_state.result = res
                            st.success(f"✅ 已识别 {len(regions)} 个区域、{len(res['systems'])} 个系统，"
                                       f"主要矛盾：**{res['principal']['system']}**")
            # —— 结构化（Excel / CSV）——
            else:
                try:
                    df = read_uploaded_file(uploaded)
                except Exception as e:
                    st.error(f"读取文件失败：{e}（读 .xls 需安装 xlrd，.xlsx 需 openpyxl：`pip install xlrd openpyxl`）")
                    df = None
                if df is not None:
                    region_cols = [c for c in df.columns if c in REGIONS]
                    st.caption(f"已读取 {len(df)} 条记录，识别到 {len(region_cols)} 个区域列。")
                    if not region_cols:
                        st.warning("未在文件中找到已知的 HRA 区域列，请确认是惠斯安普 HRA 导出格式。")
                    else:
                        opts = [record_label(df, i) for i in range(len(df))]
                        sel = st.selectbox("选择要解读的记录", opts, key="rec_sel")
                        if st.button("✨ 解读选中记录", type="primary", use_container_width=True, key="gen_file"):
                            idx = opts.index(sel)
                            profile, regions = df_row_to_data(df, idx)
                            if not regions:
                                st.error("该记录未解析出区域检测值。")
                            else:
                                res = analyze(regions, profile, st.session_state.get("q", {}))
                                st.session_state.result = res
                                st.success(f"✅ 记录「{sel}」已识别 {len(regions)} 个区域、"
                                           f"{len(res['systems'])} 个系统，主要矛盾：**{res['principal']['system']}**")

    with tab_photo:
        uploaded_list = st.file_uploader(
            "📷 拍照上传 / 图片识别（支持多张 jpg / png，可一次多选整份报告）",
            type=["jpg", "jpeg", "png"], accept_multiple_files=True)
        if uploaded_list:
            for _u in uploaded_list:
                st.image(_u.getvalue(), caption=_u.name, width=300)
            if st.button("✨ 识别并解读（多页合并）", type="primary", use_container_width=True, key="gen_photo"):
                _texts = []
                with st.spinner(f"正在 OCR 识别 {len(uploaded_list)} 张图片…"):
                    for _u in uploaded_list:
                        _t, _e = ocr_image_to_text(_u.getvalue())
                        if _e:
                            st.warning(f"⚠️ {_u.name} 识别失败：{_e}")
                        elif _t and _t.strip():
                            _texts.append(_t)
                _merged = "\n\n".join(_texts)
                if not _merged.strip():
                    st.error("所有图片都未能识别出文字。建议：用扫描 App（如手机相机'文档扫描'）拍清晰、"
                             "平整的报告，或改用「📎 上传文件」上传带文字层的 PDF / Word。")
                else:
                    profile, regions = extract_all(_merged)
                    if not regions:
                        st.error("已从图片识别出文字，但未能匹配到已知 HRA 区域。可能 OCR 把区域名识别错了。"
                                 "建议：光线均匀、正对拍平；或改用「📎 上传文件」上传带文字层的 PDF / Word。")
                    else:
                        res = analyze(regions, profile, st.session_state.get("q", {}))
                        st.session_state.result = res
                        st.success(f"✅ 共处理 {len(uploaded_list)} 张图，匹配到 {len(regions)} 个区域、"
                                   f"{len(res['systems'])} 个系统，主要矛盾：**{res['principal']['system']}**")

    # 展示结果
    res = st.session_state.result
    if res:
        st.divider()
        # ============ 首页概览：一屏看懂 ============
        hi = res["health_index"]
        st.markdown("## 🏠 一屏看懂（建议先看这里）")

        c1, c2 = st.columns([1, 2])
        with c1:
            st.markdown(
                f'<div style="text-align:center;background:{hi["color"]};'
                f'border-radius:14px;padding:18px 8px;color:#fff;">'
                f'<div style="font-size:13px;opacity:.92;">综合健康指数</div>'
                f'<div style="font-size:48px;font-weight:700;line-height:1.05;">{hi["score"]}</div>'
                f'<div style="font-size:15px;font-weight:600;">{hi["grade"]}</div></div>',
                unsafe_allow_html=True)
        with c2:
            st.markdown(f"**一句话解读**：{hi['desc']}")
            if res.get("goal"):
                st.markdown(f"🎯 **首要目标**：{res['goal']}")
            if res.get("goal_note"):
                st.markdown(f"💡 {res['goal_note']}")
            pr = res["principal"]
            if pr:
                st.markdown(f"🔴 **主要矛盾**：{pr['system']}（{pr['risk']}）")
            else:
                st.markdown("✅ 暂无突出矛盾系统。")
            ba = res.get("bio_age")
            if ba:
                st.markdown(f"🧬 **生理年龄**：约 {ba['bio']} 岁（实际 {ba['age']} 岁，"
                            f"比实际{ba['tag']} {abs(ba['offset'])} 岁）")

        # 综合健康指数分段图例
        _legend = [
            ("健康", "≥ 85", "#2ecc71", "各系统功能协调"),
            ("良好", "70–84", "#3498db", "少数系统有苗头，早干预收益大"),
            ("亚健康", "55–69", "#f1c40f", "已亮黄灯，现在干预还来得及"),
            ("警戒", "< 55", "#e74c3c", "多系统偏差明显，建议尽快就医"),
        ]
        _lg = "".join(
            f'<span style="display:inline-block;margin:3px 12px 3px 0;font-size:12px;">'
            f'<span style="display:inline-block;width:10px;height:10px;border-radius:50%;'
            f'background:{c};margin-right:5px;vertical-align:middle;"></span>'
            f'{g} <b>{r}</b> · {d}</span>'
            for g, r, c, d in _legend)
        st.markdown(f'<div style="margin:8px 0 2px;color:#555;line-height:1.8;">'
                    f'综合健康指数分段：{_lg}</div>', unsafe_allow_html=True)

        st.markdown("### ⚡ 60 秒行动（今天就能做）")
        if res["top3"]:
            for i, (dim, line) in enumerate(res["top3"], 1):
                st.markdown(f"{i}. **[{dim}]** {line}")
        else:
            st.markdown("- 保持现有良好习惯即可。")

        # 健康 OKR
        okr = res["okr"]
        st.markdown("### 🎯 健康 OKR")
        st.markdown(f"**O（目标）**：{okr['o']}")
        for k in okr["krs"]:
            st.markdown(f"- **KR**：{k}")
        # 治未病窗口
        tp = res["tipping"]
        if tp:
            st.warning("🌱 **治未病窗口（质变前最后机会）**："
                       + "、".join(f"{s}（{r}，最大偏离 {m}）" for s, m, r in tp)
                       + " 已接近高风险阈值——此刻干预成本最低，拖下去将进入『质变』。")

        rc = _radar_chart(res["systems"])
        a1, a2 = st.columns([1, 1])
        with a1:
            if rc:
                st.image(rc, use_column_width=True)
            else:
                st.caption("区域不足 3 个系统，暂不可绘制雷达图。")
        with a2:
            st.markdown("#### 🟢 健康资产（保护好）")
            if res["assets"]["assets"]:
                st.markdown("、".join(f"`{x}`" for x in res["assets"]["assets"]))
            else:
                st.markdown("（暂无完全正常的系统，先稳住主要矛盾）")
            st.markdown("#### 🔴 健康负债（优先偿还）")
            if res["assets"]["liabilities"]:
                st.markdown("、".join(f"`{x}`（{r}）" for x, r in res["assets"]["liabilities"]))
            else:
                st.markdown("（当前无异常负债，继续保持）")

        st.markdown("### 🧠 认知提醒 ＆ 🛡️ 抗风险提醒")
        st.caption("人常因『现在没症状』（现状偏差）和『感觉良好』（过度自信）而低估慢性风险——"
                   "报告测的是功能偏差，未必有感觉，请以数据为准。")
        st.caption("健康要『抗风险』——适度良性压力（快走、规律作息）让你在波动中变强；"
                   "要躲开的是脆性行为与『小概率高危害』的冒险（熬夜、猛药、极端节食）。")

        st.divider()
        # ============ 正式报告正文 ============
        st.markdown(f"## 一、总体概览\n{res['overview']}")

        # 系统风险条形图（横向，系统名正排可读）
        st.markdown("### 📊 各系统风险指数（0–100）")
        st.caption("分数按功能偏离强度归一化到 0–100（**不是**百分制累加）："
                   "**≥70 高风险** ｜ **40–69 中风险** ｜ **20–39 低风险** ｜ **<20 正常**。"
                   "右侧括号内为原始「累计偏离强度」（所有区域绝对偏离值之和，可超过 100，如骨骼系统常达 1000+）。")
        _systems = res["systems"]
        _max = max((s["score"] for s in _systems), default=1) or 1
        _color = {"高风险": "#e74c3c", "中风险": "#e67e22",
                  "低风险": "#f1c40f", "正常": "#2ecc71"}
        _bars = []
        for s in _systems:
            pct = int(s["score"] / _max * 100) if _max else 0
            c = _color.get(s["risk"], "#3498db")
            _bars.append(
                f'<div style="display:flex;align-items:center;margin:5px 0;">'
                f'<div style="width:104px;font-size:13px;text-align:right;'
                f'padding-right:8px;white-space:nowrap;">{s["system"]}</div>'
                f'<div style="flex:1;background:#eceff1;border-radius:4px;height:18px;">'
                f'<div style="width:{pct}%;background:{c};height:18px;border-radius:4px;"></div></div>'
                f'<div style="width:120px;font-size:13px;padding-left:8px;">{s["score"]}'
                f' <span style="color:#888;">(偏离{s["raw"]})</span></div></div>'
            )
        st.markdown('<div style="font-family:sans-serif;max-width:760px;">'
                    + "".join(_bars) + "</div>", unsafe_allow_html=True)

        # 逐系统
        st.markdown("## 二、逐系统局部解读")
        for s in res["systems"]:
            icon = s["meta"].get("icon", "")
            with st.expander(f"{icon} {s['system']} —— 风险：**{s['risk']}**"
                             f"（异常 {s['n_abn']}/{s['n_total']}，最大偏离 {s['maxabs']}）",
                             expanded=(s["risk"] in ("高风险", "中风险"))):
                st.markdown(f"**当前状态**：{s['findings']}")
                if s["abn"]:
                    st.markdown("**主要偏差区域**：" +
                                "、".join(f"`{r} = {v}`" for r, v in s["items"][:5]))

        # 综合研判（矛盾论）
        st.markdown("## 三、综合研判：主要矛盾与主要方面")
        pr = res["principal"]
        if pr:
            st.markdown(f"🔴 **主要矛盾**：**{pr['system']}**（风险 {pr['risk']}，"
                        f"最大偏离 {pr['maxabs']}）——各系统中功能偏差最突出。")
            if res["primary"]:
                r, v = res["primary"]
                st.markdown(f"🎯 **主要方面**：该系统最关键区域 **{r}**（检测值 {v}），"
                            f"应优先干预。")
            if res["secondary"]:
                st.markdown("🔸 **次要方面**：" +
                            "、".join(f"{r}（{v}）" for r, v in res["secondary"]))
            others = [s for s in res["systems"] if s is not pr and s["risk"] != "正常"][:2]
            if others:
                st.markdown("🔹 **次要矛盾**：" +
                            "、".join(f"{s['system']}（{s['risk']}）" for s in others))

        # 系统联动提示
        if res["links"]:
            st.markdown("## 🔗 系统联动提示")
            st.markdown("> 身体是耦合系统，单点异常常牵动一串。下面是被同时点亮的『连锁链』：")
            for chain, txt in res["links"]:
                st.markdown(f"- **{chain}**：{txt}")

        # 行动优先级 Top3
        st.markdown("## 四、🎯 行动优先级 Top 3（先做好这三件）")
        if res["top3"]:
            for i, (dim, line) in enumerate(res["top3"], 1):
                st.markdown(f"{i}. **[{dim}]** {line}")
            st.info("以上三项针对你的主要矛盾系统，性价比最高，建议本周就开始执行。")
        else:
            st.markdown("暂无突出矛盾，保持现有良好生活习惯即可。")

        # 行动打卡（行为闭环 / 成功日记）
        st.markdown("## 📝 行动打卡（行为大于知识）")
        _chk = []
        for dim, line in res["top3"]:
            _chk.append(f"【{dim}】{line}")
        for d in ["饮食", "运动", "睡眠", "情绪管理", "生活方式"]:
            _first = res["six"].get(d, [""])[0] if res["six"].get(d) else None
            if _first:
                _chk.append(f"【{d}】{_first}")
        for i, item in enumerate(_chk[:8]):
            st.checkbox(item, key=f"chk_{i}")
        st.caption("✅ 每完成一项打勾，并写一句『成功日记』：今天我做到了___。坚持 21 天，习惯变资产。")

        # 逆向清单
        st.markdown("## 🚫 逆向清单：先别做蠢事")
        st.markdown("> 进步最快的方式，是先停止做蠢事。把下面这些『别做』刻进习惯：")
        for it in res["invert"]:
            st.markdown(f"- {it}")

        # 六维综合干预建议（统一，不逐系统罗列）
        st.markdown("## 五、六维综合干预建议（统一方案）")
        dim_icon = {"饮食": "🥗", "运动": "🏃", "睡眠": "🌙",
                    "用药": "💊", "情绪管理": "🧘", "生活方式": "💡"}
        for d in ["饮食", "运动", "睡眠", "用药", "情绪管理", "生活方式"]:
            with st.expander(f"{dim_icon.get(d, '')} {d}", expanded=False):
                for line in res["six"][d]:
                    st.markdown(f"- {line}")
                # 饮食维度补充明细
                if d == "饮食":
                    dd = res["diet_detail"]
                    if dd.get("食物"):
                        st.markdown("**🍽️ 推荐食物（及作用）**")
                        for it in dd["食物"]:
                            st.markdown(f"&nbsp;&nbsp;- {it}")
                    if dd.get("茶饮"):
                        st.markdown("**🍵 推荐茶饮**")
                        for it in dd["茶饮"]:
                            st.markdown(f"&nbsp;&nbsp;- {it}")
                    if dd.get("禁忌"):
                        st.markdown("**🚫 饮食禁忌**")
                        for it in dd["禁忌"]:
                            st.markdown(f"&nbsp;&nbsp;- {it}")
                    if dd.get("习惯"):
                        st.markdown("**⏰ 用餐习惯**")
                        for it in dd["习惯"]:
                            st.markdown(f"&nbsp;&nbsp;- {it}")
            st.divider()

        # 就医建议
        st.markdown("## 六、🏥 就医建议（必要时及时就诊）")
        st.info("📎 就诊时请出示本报告（可下载 .md / 基线 JSON），便于医生快速了解你的 HRA 风险分布。")
        if res["medical"]:
            med_df = {"系统": [m[0] for m in res["medical"]],
                      "建议科室": [m[1] for m in res["medical"]],
                      "推荐检查项目": [m[2] for m in res["medical"]]}
            st.table(med_df)
        else:
            st.markdown("当前无系统显著异常，常规年度体检即可。")

        # 用药建议
        st.markdown("## 七、💊 用药建议（务必遵医嘱）")
        st.warning("⚠️ **重要提醒**：以下仅为科普参考。**任何药物都必须在执业医师明确诊断后、"
                   "遵医嘱使用；不可自行购药、不可随意加量/减量/停药，以免病情反弹或出现戒断、耐药等风险。**")
        if res["medication"]:
            for sysname, txt in res["medication"]:
                st.markdown(f"- **{sysname}**：{txt}")
        else:
            st.markdown("当前无明确用药指征，请勿自行服药。")

        # 身体状态趋势预警 + 复测跟踪
        st.markdown("## 八、⚠️ 身体状态趋势预警（若不干预）")
        if res["trend"]:
            for sysname, txt in res["trend"]:
                st.markdown(f"- **{sysname}**：{txt}")
        else:
            st.markdown("当前风险尚低，保持即可。")
        st.info(f"📅 **复测与跟踪（健康储蓄账户）**：{res['retest']}")
        _baseline = {
            "date": datetime.date.today().isoformat(),
            "profile": res["profile"],
            "health_index": res["health_index"]["score"],
            "systems": [{"system": s["system"], "score": s["score"],
                         "maxabs": s["maxabs"], "risk": s["risk"]} for s in res["systems"]],
        }
        st.download_button("💾 保存本次为基线（下次复测对比）",
                           json.dumps(_baseline, ensure_ascii=False, indent=2),
                           "HRA基线.json", "application/json")

        # 健康宣教科普
        st.markdown("## 九、📖 健康宣教科普")
        if res["edu"]:
            for sysname, txt in res["edu"]:
                st.markdown(f"- **{sysname}**：{txt}")
        else:
            st.markdown("保持良好生活方式是最好的预防。")

        # 健康轨迹
        st.markdown("## 📈 健康轨迹")
        st.info("上传你之前『保存的基线 JSON』（可多选），与本次叠加看趋势——健康是长期复利，关键看曲线方向。")
        _hist = st.file_uploader("📎 上传历史基线 JSON（可多选）",
                                 type=["json"], accept_multiple_files=True, key="traj_up")
        _traj = []
        if _hist:
            for _fh in _hist:
                try:
                    _traj.append(json.loads(_fh.getvalue().decode("utf-8")))
                except Exception:
                    st.warning(f"无法解析 {_fh.name}")
        _traj.append({
            "date": datetime.date.today().isoformat(),
            "profile": res["profile"],
            "health_index": res["health_index"]["score"],
            "systems": [{"system": s["system"], "score": s["score"],
                         "maxabs": s["maxabs"], "risk": s["risk"]} for s in res["systems"]],
        })
        if len(_traj) >= 2:
            _draw_trajectory(_traj)
        else:
            st.caption("至少 2 次数据（历史 + 本次）才能画趋势曲线；先保存本次基线，下次复测时上传对比。")

        md = to_markdown(res)
        st.download_button("📥 下载完整报告 (.md)", md, "HRA综合解读报告.md", "text/markdown")
        # —— 下载 PDF 版报告：双击即排版好的中文 PDF，无需纠结 .md 怎么看 ——
        try:
            _radar_pdf = _radar_chart(res["systems"])
            _pdf_bytes = _md_to_pdf(md, _radar_pdf)
            st.download_button("📄 下载 PDF 报告（排版好，双击即用）", _pdf_bytes,
                               "HRA综合解读报告.pdf", "application/pdf")
        except Exception as _e:
            st.caption(f"⚠️ PDF 生成暂不可用：{_e}")
        st.caption("⚠️ 本报告由 AI 辅助生成，仅作健康科普与自我管理参考，不构成诊断或治疗建议。"
                   "异常指标请务必咨询执业医师。")

# —— 免装部署：让不懂 Python 的朋友也能直接打开网页 ——
with st.expander("🚀 免装部署：如何发给不懂 Python 的朋友直接用？"):
    st.markdown("""
**方案 A（推荐 · 免费 · 永久在线）—— Streamlit Cloud**
1. 把本文件夹推到你的 GitHub 仓库（依赖文件 `requirements.txt` 已备好，部署时会自动安装）；
2. 打开 https://share.streamlit.io ，用 GitHub 登录，选该仓库、入口填 `app.py`；
3. 点 **Deploy**，几分钟后得到一个 `https://xxx.streamlit.app` 链接——**发给任何人点开即用**，无需装任何东西。

**方案 B（本地临时分享）—— ngrok**
本机网页跑起来后，另开一个终端：`ngrok http 8501`，把生成的公网地址发给朋友即可。

**方案 C（发给 Windows 朋友双击用）—— 打包 exe**
本机终端运行：`pip install pyinstaller && pyinstaller --onefile app.py`，
生成的 `dist/app.exe` 直接发给朋友双击即可（首次启动稍慢）。

> 你现在是在自己电脑**本地运行**（地址 `http://localhost:8501`），只有你自己能打开。要让别人用，选上面任一方案即可。
""")
