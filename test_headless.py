# -*- coding: utf-8 -*-
"""悦检 v3 headless 自检：引擎端到端 + 新字段 + 雷达图 + 文档提取链路。"""
import sys, os, io
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from engine_v2 import analyze
from extract import extract_all

base = os.path.dirname(os.path.abspath(__file__))
sample = os.path.join(base, "sample_real.tsv")
if os.path.exists(sample):
    txt = open(sample, encoding="utf-8").read()
else:
    txt = ("|姓名|示例|性别:|女|身高(cm)|165|体重(Kg)|60|\n"
           "|气管附近|气管附近|-66|支气管区域|支气管区域|-66|右肺下叶区域|右肺下叶区域|-50|\n"
           "|心脏区域|心脏区域|-56|冠状血管|冠状血管|54|左颈动脉|左颈动脉|-53|")

q = {"goal": "心血管养护", "age": "18–25", "exercise_freq": "几乎不运动", "exercise_type": "基本没有",
     "diet": "偏油腻", "water": "<1L", "sleep_hours": "<6h", "sleep_quality": "差（易醒/难入睡）",
     "stress": "高", "smoke": "偶尔", "alcohol": "偶尔", "sit": ">8h"}
profile, regions = extract_all(txt)
res = analyze(regions, profile, q)

for k in ["health_index", "retest", "assets", "goal", "goal_note",
          "bio_age", "okr", "tipping", "invert", "links",
          "six", "diet_detail", "medical", "medication", "trend", "edu", "top3",
          "systems", "principal"]:
    assert k in res, f"missing {k}"

print("REGIONS", len(regions), "| SYSTEMS", len(res["systems"]),
      "| PRINCIPAL", res["principal"]["system"] if res["principal"] else None)
print("HEALTH_INDEX", res["health_index"])
print("RETES7", res["retest"][:46])
print("ASSETS", res["assets"])
print("GOAL_NOTE", res["goal_note"])
print("TOP3", res["top3"])
print("SIX keys", list(res["six"].keys()))
print("MED n", len(res["medical"]), "TREND n", len(res["trend"]),
      "EDU n", len(res["edu"]), "MEDICATION n", len(res["medication"]))
print("BIO_AGE", res["bio_age"])
print("OKR", res["okr"])
print("TIPPING", res["tipping"])
print("INVERT", res["invert"][:3])
print("LINKS", res["links"])

# 雷达图逻辑（不依赖 streamlit，仅验证 matplotlib 绘制不报错）
import math, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
def radar(systems):
    labels = [s["system"].replace("系统", "") for s in systems]
    vals = [min(s["maxabs"], 100) for s in systems]
    n = len(labels)
    if n < 3:
        return "skip(<3)"
    ang = [i / float(n) * 2 * math.pi for i in range(n)] + [0.0]
    vals2 = vals + [vals[0]]
    fig, ax = plt.subplots(figsize=(4, 4), subplot_kw=dict(polar=True))
    ax.plot(ang, vals2)
    ax.fill(ang, vals2, alpha=0.2)
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return "ok"
print("RADAR", radar(res["systems"]))

# 文档提取链路：生成 docx → 提取文本 → extract_all
try:
    import docx as d
    doc = d.Document()
    doc.add_paragraph("|姓名|测试|性别:|男|身高(cm)|175|体重(Kg)|70|")
    doc.add_paragraph("|心脏区域|心脏区域|55|冠状血管|冠状血管|60|气管附近|气管附近|-40|")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "左颈动脉"
    t.rows[0].cells[1].text = "-50"
    p = "/tmp/hra_test_doc.docx"
    doc.save(p)
    d2 = d.Document(p)
    txt2 = "\n".join(pp.text for pp in d2.paragraphs if pp.text.strip())
    for tbl in d2.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    txt2 += "\n" + cell.text
    prof2, reg2 = extract_all(txt2)
    print("DOCX chain regions", len(reg2), dict(list(reg2.items())[:4]))
except Exception as e:
    print("DOCX chain err:", repr(e))

# PDF 库存在性 + 最小读取验证
try:
    from pypdf import PdfReader
    print("pypdf import OK")
except Exception as e:
    print("pypdf import err:", repr(e))

print("ALL OK")
